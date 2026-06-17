#!/usr/bin/env python3
"""Build the editable Word version (v2) of the Maths Prompt Studio from data/prompts.js."""
import json, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SIG = "Indrajeet Yadav"
BLUE = RGBColor(0x2b,0x3f,0x6b); GOLD = RGBColor(0xb1,0x82,0x2c); VIOLET = RGBColor(0x6b,0x4a,0xa0)
GREY = RGBColor(0x55,0x52,0x4b); INK = RGBColor(0x23,0x26,0x2d); GREEN = RGBColor(0x2f,0x7d,0x56)

src = open(os.path.join(ROOT, "data", "prompts.js"), encoding="utf-8").read()
DATA = json.loads(src[src.index("window.PROMPT_DATA =")+len("window.PROMPT_DATA ="): src.rindex(";")])
CATS = DATA["categories"]
TOTAL = sum(len(c["prompts"]) for c in CATS)
GROUP_ORDER = ['Solving & Checking','Practice & Assessment','Teaching Materials','Writing & Content','Engagement','Support','Teacher Productivity']
GROUPS = [g for g in GROUP_ORDER if any(c.get("group")==g for c in CATS)]
for c in CATS:
    if c.get("group") not in GROUPS: GROUPS.append(c.get("group") or "More")

def clean(s): return re.sub(r"[^\x00-\x7f]", "", str(s or "")).strip()

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8); s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

def para(text="", size=10.5, bold=False, italic=False, color=None, align=None, sa=6, sb=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(sb)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

# cover
for _ in range(4): doc.add_paragraph()
para("THE MATHS PROMPT STUDIO", 12, True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
para(f"{TOTAL} AI Prompts for Mathematics Teachers", 28, True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)
para("Copy-paste prompts that turn any maths problem into handwritten solutions, question papers, worksheets, DPPs, formula sheets, books and more - with a complete beginner's guide. Free forever.", 12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=22)
para(f"{TOTAL} prompts   .   {len(CATS)} categories   .   {len(CATS[0]['prompts'])} handwritten art styles", 11, True, color=VIOLET, align=WD_ALIGN_PARAGRAPH.CENTER, sa=28)
para("Created and authored by", 10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
para(SIG, 22, True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# how to (beginner)
para("Brand new to AI? Start here.", 20, True, color=BLUE, sa=8)
para("AI tools (ChatGPT, Gemini, Claude, Copilot) are free websites where you type a request - or attach a photo - and get a written answer back, like texting a brilliant assistant. No installing, no coding.", 11, color=GREY, sa=10)
para("The 4 steps", 13, True, color=VIOLET, sa=6)
steps = [
 ("1","Open a free AI tool."," ChatGPT (chatgpt.com) or Gemini (gemini.google.com) are easiest. Sign in with Google - free."),
 ("2","Attach a photo (if the prompt needs one)."," In the typing bar tap the + or paperclip on the left, choose Camera or Photos, and pick a clear picture of ONE question."),
 ("3","Copy a prompt from this book and paste it."," Long-press the typing bar and tap Paste (phone), or Ctrl/Cmd + V (computer)."),
 ("4","Fill the [BRACKETS] and send."," Replace [CLASS], [BOARD], [TOPIC] with your details, press Send. Download or print the result."),
]
for n,b,rest in steps:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7)
    r = p.add_run(n+".  "); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=GOLD
    r = p.add_run(b); r.bold=True; r.font.size=Pt(11)
    r = p.add_run(rest); r.font.size=Pt(11); r.font.color.rgb=GREY
para("What works on a free plan? Text tasks (solutions, papers, worksheets, notes, books) work on ANY free AI. Picture tasks (the handwritten art pages) need an image-making AI and free plans limit images per day - each prompt tells you which.", 10, italic=True, color=VIOLET, sb=4, sa=6)
para(f"Every prompt asks the AI to sign the work as '{SIG}', so your name travels with everything you create. Please keep it when you share.", 10, italic=True, color=GREY)
doc.add_page_break()

# contents
para("Contents", 20, True, color=BLUE, sa=8)
n = 0
for g in GROUPS:
    para(g, 13, True, color=VIOLET, sb=8, sa=2)
    for c in [c for c in CATS if c.get("group")==g]:
        n += 1
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"   {n}.  {clean(c['categoryTitle'])}"); r.bold=True; r.font.size=Pt(11)
        r = p.add_run(f"   ({len(c['prompts'])})"); r.font.size=Pt(9.5); r.font.color.rgb=GREY
doc.add_page_break()

def rel_label(p):
    return "Makes images (needs an image AI)" if p.get("makesImage") else ("Attach a photo first" if p.get("needsImage") else "Works on any free AI")

# sections by group
n = 0
for g in GROUPS:
    para(g.upper(), 16, True, color=VIOLET, align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=10)
    for c in [c for c in CATS if c.get("group")==g]:
        n += 1
        para(f"Section {n}", 9, True, color=GOLD, sa=0)
        para(clean(c["categoryTitle"]), 19, True, color=BLUE, sa=2)
        para(clean(c["categoryBlurb"]), 10, italic=True, color=GREY, sa=10)
        for pi, p in enumerate(c["prompts"], 1):
            para(f"{n}.{pi}   {clean(p['title'])}", 13, True, color=VIOLET, sb=8, sa=2)
            para(f"[{rel_label(p)}]   Best tool: {clean(p.get('bestTool','Any AI chat'))}", 9, True, color=GREEN if not (p.get('makesImage') or p.get('needsImage')) else GOLD, sa=2)
            wp = doc.add_paragraph(); wp.paragraph_format.space_after = Pt(1)
            r = wp.add_run("What you get: "); r.bold=True; r.font.size=Pt(9.5)
            r = wp.add_run(clean(p["whatYouGet"])); r.font.size=Pt(9.5); r.font.color.rgb=GREY
            eff = p.get("effectiveUsage") or []
            if eff:
                hp = doc.add_paragraph(); hp.paragraph_format.space_after = Pt(1)
                r = hp.add_run("How to use: "); r.bold=True; r.font.size=Pt(9.5)
                r = hp.add_run("  ".join(f"({i+1}) {clean(s)}" for i,s in enumerate(eff))); r.font.size=Pt(9.5); r.font.color.rgb=GREY
            if p.get("commonFix"):
                fp = doc.add_paragraph(); fp.paragraph_format.space_after = Pt(4)
                r = fp.add_run("If it is wrong, reply: "); r.bold=True; r.font.size=Pt(9.5)
                r = fp.add_run(clean(p["commonFix"])); r.font.size=Pt(9.5); r.font.color.rgb=GREY
            para("COPY THIS PROMPT", 8, True, color=GOLD, sa=2)
            tbl = doc.add_table(rows=1, cols=1); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0,0); pr = cell._tc.get_or_add_tcPr()
            sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "F3ECDD"); pr.append(sh)
            cp = cell.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
            rr = cp.add_run(clean(p["promptText"])); rr.font.name="Consolas"; rr.font.size=Pt(8); rr.font.color.rgb=INK
        doc.add_page_break()

ftr = doc.sections[0].footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = ftr.add_run(f"Maths Prompt Studio  .  Created & authored by {SIG}  .  Free forever"); fr.font.size=Pt(8); fr.font.color.rgb=GREY

os.makedirs(os.path.join(ROOT, "downloads"), exist_ok=True)
out = os.path.join(ROOT, "downloads", "Maths-Prompt-Studio-by-Indrajeet-Yadav.docx")
doc.save(out)
print("wrote", out, f"({TOTAL} prompts)")
